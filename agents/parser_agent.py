"""
ParserAgent v3.0: Extracts richly structured information from resume files.
Supports PDF, DOCX, and plain text. Extracts contact info, experience,
education, skills, projects, certifications, and location.
"""

import re
import io
from typing import Dict, Any, List
from agents.knowledge_base import KNOWN_SKILLS


class ParserAgent:
    """
    Agent responsible for parsing resume files into structured data.
    Uses rule-based NLP + regex for speed with LLM-style output format.
    """

    SECTION_HEADERS = {
        "experience":     ["experience", "work history", "employment", "career", "professional background", "work experience"],
        "education":      ["education", "academic", "qualification", "degree", "university", "academic background"],
        "skills":         ["skills", "technologies", "tech stack", "competencies", "expertise", "tools", "technical skills", "core competencies"],
        "projects":       ["projects", "portfolio", "work samples", "key projects", "personal projects", "side projects"],
        "certifications": ["certifications", "certificates", "credentials", "licenses", "awards", "achievements"],
        "summary":        ["summary", "objective", "profile", "about", "overview", "professional summary"],
    }

    EMAIL_RE    = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    PHONE_RE    = re.compile(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]')
    LINKEDIN_RE = re.compile(r'linkedin\.com/in/[\w\-]+')
    GITHUB_RE   = re.compile(r'github\.com/[\w\-]+')
    YEARS_RE    = re.compile(r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)?', re.IGNORECASE)
    LOCATION_RE = re.compile(r'\b([A-Z][a-zA-Z\s]+,\s*(?:[A-Z]{2}|[A-Z][a-z]+))\b')

    # Date range for experience parsing
    DATE_RANGE_RE = re.compile(
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*)?(\d{4})\s*[-–—]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*)?(\d{4}|[Pp]resent|[Cc]urrent)',
        re.IGNORECASE
    )

    async def parse(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Main parse entry point — routes by file type."""
        if filename.lower().endswith('.pdf'):
            text = self._parse_pdf(content)
        elif filename.lower().endswith('.docx'):
            text = self._parse_docx(content)
        else:
            text = content.decode('utf-8', errors='ignore')
        return self._extract_structure(text, filename)

    def _parse_pdf(self, content: bytes) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                return "\n".join(pages)
        except ImportError:
            pass
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return content.decode('utf-8', errors='ignore')

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = [para.text for para in doc.paragraphs]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except ImportError:
            return content.decode('utf-8', errors='ignore')

    def _extract_structure(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract all structured fields from raw text."""
        lines = [l.strip() for l in text.split('\n') if l.strip()]

        name     = self._extract_name(lines)
        email    = self._find_first(self.EMAIL_RE, text)
        phone    = self._find_first(self.PHONE_RE, text)
        linkedin = self._find_first(self.LINKEDIN_RE, text)
        github   = self._find_first(self.GITHUB_RE, text)
        location = self._extract_location(text)

        sections = self._split_sections(text)

        skill_text = (
            sections.get("skills", "") + "\n" +
            sections.get("experience", "") + "\n" +
            text
        )
        raw_skills    = self._extract_skills_from_text(skill_text)
        years_exp     = self._extract_years(text)
        education     = self._extract_education(sections.get("education", text))
        experience    = self._extract_experience(sections.get("experience", ""))
        certifications = self._extract_certifications(sections.get("certifications", ""))
        projects      = self._extract_projects(sections.get("projects", ""))
        summary       = sections.get("summary", lines[0] if lines else "")[:400]

        return {
            "name":              name,
            "email":             email,
            "phone":             phone,
            "linkedin":          linkedin,
            "github":            github,
            "location":          location,
            "summary":           summary,
            "raw_text":          text,
            "raw_skills":        raw_skills,
            "years_of_experience": years_exp,
            "education":         education,
            "experience":        experience,
            "certifications":    certifications,
            "projects":          projects,
            "sections_detected": list(sections.keys()),
            "raw_text_length":   len(text),
            "filename":          filename,
            "status":            "parsed",
        }

    # ── Extractors ────────────────────────────────────────────────────────────

    def _extract_name(self, lines: List[str]) -> str:
        for line in lines[:6]:
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(re.match(r"^[A-Z][a-zA-Z'\-]+$", w) for w in words):
                    return line
        return lines[0] if lines else "Unknown"

    def _find_first(self, pattern, text: str) -> str:
        match = pattern.search(text)
        return match.group(0) if match else ""

    def _extract_location(self, text: str) -> str:
        """Try to extract city/state or city/country from text."""
        # Common patterns: "San Francisco, CA" or "Bangalore, India"
        city_state = re.search(
            r'\b([A-Z][a-zA-Z\s]{2,20}),\s*([A-Z]{2}|[A-Z][a-z]{3,15})\b',
            text
        )
        if city_state:
            return city_state.group(0)
        return ""

    def _split_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        current_section = "other"
        current_lines: List[str] = []

        for line in text.split('\n'):
            lower = line.lower().strip()
            matched_section = None
            for section_key, keywords in self.SECTION_HEADERS.items():
                if any(lower == kw or lower.startswith(kw + ':') or lower == kw + 's' for kw in keywords):
                    matched_section = section_key
                    break
            if matched_section:
                if current_lines:
                    sections[current_section] = "\n".join(current_lines)
                current_section = matched_section
                current_lines = []
            else:
                current_lines.append(line)

        if current_lines:
            sections[current_section] = "\n".join(current_lines)
        return sections

    def _extract_skills_from_text(self, text: str) -> List[str]:
        found = []
        text_lower = text.lower()
        for skill in KNOWN_SKILLS:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found.append(skill)
        return found

    def _extract_years(self, text: str) -> int:
        matches = self.YEARS_RE.findall(text)
        if matches:
            return max(int(m) for m in matches)
        date_ranges = re.findall(
            r'(20\d\d|19\d\d)\s*[-–]\s*(20\d\d|19\d\d|present|current)',
            text, re.IGNORECASE
        )
        if date_ranges:
            import datetime
            current_year = datetime.datetime.now().year
            total = 0
            for start, end in date_ranges:
                s = int(start)
                e = current_year if end.lower() in ('present', 'current') else int(end)
                total += max(0, e - s)
            return min(total, 20)
        return 0

    def _extract_education(self, text: str) -> List[Dict]:
        degrees = [
            "Bachelor", "Master", "PhD", "Ph.D", "B.Tech", "M.Tech",
            "B.E", "M.E", "B.Sc", "M.Sc", "MBA", "BCA", "MCA",
            "Associate", "Diploma", "B.S", "M.S", "B.A", "M.A"
        ]
        results = []
        for line in text.split('\n'):
            for degree in degrees:
                if degree.lower() in line.lower():
                    results.append({"degree": degree, "line": line.strip()})
                    break
        return results[:4]

    def _extract_experience(self, text: str) -> List[Dict]:
        """Parse work experience entries from the experience section."""
        entries = []
        if not text.strip():
            return entries

        # Split by date ranges or company patterns
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        current: Dict[str, Any] = {}
        responsibilities: List[str] = []

        for line in lines:
            has_date = bool(self.DATE_RANGE_RE.search(line))
            is_bullet = line.startswith(('-', '•', '*', '·', '▪')) or (len(line) > 2 and line[0].isdigit() and line[1] == '.')

            if has_date or (re.search(r'\b(at|@|,)\b', line, re.I) and len(line.split()) <= 8):
                if current:
                    current["responsibilities"] = responsibilities[:6]
                    entries.append(current)
                    responsibilities = []
                # Try to parse role and company
                parts = re.split(r'\s*(?:at|@|—|–|-)\s*', line, maxsplit=1)
                role = parts[0].strip() if parts else line
                company = parts[1].strip() if len(parts) > 1 else ""
                # Strip date from company if present
                company = self.DATE_RANGE_RE.sub('', company).strip().strip(',').strip()
                duration_match = self.DATE_RANGE_RE.search(line)
                duration = duration_match.group(0) if duration_match else ""

                current = {"role": role[:80], "company": company[:80], "duration": duration}
            elif is_bullet and current:
                clean = re.sub(r'^[-•*·▪\d\.]\s*', '', line).strip()
                if clean:
                    responsibilities.append(clean[:200])
            elif current and not is_bullet and len(line) > 10:
                # Could be a continuation or company name clarification
                if not current.get("company"):
                    current["company"] = line[:80]

        if current:
            current["responsibilities"] = responsibilities[:6]
            entries.append(current)

        return entries[:8]

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certification lines."""
        certs = []
        cert_keywords = ["certified", "certification", "certificate", "aws", "gcp", "azure",
                         "professional", "associate", "practitioner", "pmp", "scrum", "ckad",
                         "cka", "comptia", "cissp", "ccna", "google", "microsoft"]
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            lower = line.lower()
            if any(kw in lower for kw in cert_keywords) and len(line) > 5:
                clean = re.sub(r'^[-•*·▪\d\.]\s*', '', line).strip()
                if clean and clean not in certs:
                    certs.append(clean)
        return certs[:8]

    def _extract_projects(self, text: str) -> List[str]:
        """Extract project names/descriptions."""
        projects = []
        for line in text.split('\n'):
            line = line.strip()
            if not line or len(line) < 5:
                continue
            clean = re.sub(r'^[-•*·▪\d\.]\s*', '', line).strip()
            if clean and len(clean) > 10:
                projects.append(clean[:200])
        return projects[:6]
